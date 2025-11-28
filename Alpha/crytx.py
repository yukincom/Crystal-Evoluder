"""
Crystal Evoluder v1.1
Knowledge Crystallization System

"""

import os
import re
import logging
import concurrent.futures
import time
import json                                
from datetime import datetime  
from pathlib import Path
from typing import Optional, List, Dict, Tuple,Any
from contextlib import contextmanager

import requests   
from pathlib import Path
from typing import List
from bs4 import BeautifulSoup

from llama_index.core import Document, KnowledgeGraphIndex, StorageContext
from llama_index.graph_stores.neo4j import Neo4jGraphStore
from llama_index.llms.openai import OpenAI
from llama_index.core.node_parser import SimpleNodeParser
from llama_index.embeddings.huggingface import HuggingFaceEmbedding

# 追加パッケージ（オプション）
try:
    import frontmatter
    HAS_FRONTMATTER = True
except ImportError:
    HAS_FRONTMATTER = False
    print("⚠️  python-frontmatter not installed (Markdown YAML support limited)")

try:
    import ftfy
    HAS_FTFY = True
except ImportError:
    HAS_FTFY = False
    print("⚠️  ftfy not installed (text cleaning limited)")

try:
    from langchain_experimental.text_splitter import SemanticChunker 
    from langchain_openai import OpenAIEmbeddings
    HAS_SEMANTIC_CHUNKER = True
except ImportError:
    HAS_SEMANTIC_CHUNKER = False
    print("⚠️  langchain-experimental not installed (using basic chunking)")

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("⚠️  python-docx not installed (DOCX support disabled)")

try:
    import trafilatura
    HAS_TRAFILATURA = True
except ImportError:
    HAS_TRAFILATURA = False
    print("⚠️  trafilatura not installed (HTML extraction limited)")


try:
    from tqdm import tqdm
    HAS_TQDM = True

except ImportError:
    HAS_TQDM = False
    print("⚠️  tqdm not installed (progress bar disabled)")


def collect_files(target_dir: str, allowed_ext: List[str]) -> List[str]:
    files = []
    for root, _, filenames in os.walk(target_dir):
        for fname in filenames:
            if fname.lower().endswith(tuple(allowed_ext)):
                files.append(os.path.join(root, fname))
    return files

def parse_single_file(self, file_path: str) -> List[Document]:
    """単一ファイルを解析して Document の list を返す（wrapper）"""
    p = Path(file_path)
    fmt = self._detect_format(str(p))
        # crystallize は path+format を受け取って Document list を返すようにしてあるから使う
        # ここでは strict_mode は False にしておく（必要なら引数で渡す）
    return self.crystallize(str(p), format=fmt, strict_mode=False)

def parse_all_files(self, file_list: List[Path], max_workers: int = 4) -> List[Document]:
    """複数ファイルを並列処理して Document を平坦化して返す"""
    all_docs = []
    
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_file = {executor.submit(self.crystallize, str(f)): f for f in file_list}
        for future in concurrent.futures.as_completed(future_to_file):
            
            try:
                docs = future.result()
                all_docs.extend(docs)
            except Exception as e:
                self.logger.error(f"Failed to parse {future_to_file[future]}: {e}")
                self.logger.info(f"✅ All files parsed: {len(all_docs)} documents")
                return all_docs

class HierarchicalLogger:
    """階層的ログ出力"""
    
    def __init__(self, logger: logging.Logger):
        self.logger = logger
        self.indent_level = 0
    
    @contextmanager
    def section(self, title: str):
        """セクション開始"""
        self.logger.info(f"{'  ' * self.indent_level}▶ {title}")
        self.indent_level += 1
        start = time.time()
        
        try:
            yield
        finally:
            elapsed = time.time() - start
            self.indent_level -= 1
            self.logger.info(f"{'  ' * self.indent_level}✓ {title} ({elapsed:.2f}s)")
    
    def info(self, msg: str):
        self.logger.info(f"{'  ' * self.indent_level}{msg}")

class CrystalEvoluder:
    """Crystal Evoluder - Knowledge Crystallization System"""
    
    def __init__(self, config: Optional[Dict] = None, log_level: int = logging.INFO):
        self.config = config or {}
        self.crystal = None
        self.metadata = {}
        self.logger = self._setup_logger(log_level)

        self.grobid_url = self.config.get('grobid_url', 'http://localhost:8070')
        self.grobid_available = self._check_grobid()

        if self.grobid_available:
            self.logger.info(f"✅ Grobid server available at {self.grobid_url}")
        else:
            self.logger.warning("⚠️  Grobid server not available (PDF support disabled)")
        
        self.logger.info("Crystal Evoluder v1.1 initialized")

    def _setup_logger(self, level: int) -> logging.Logger:
        """ロガー設定"""
        logger = logging.getLogger('CrystalEvoluder')
        logger.setLevel(level)
        logger.handlers.clear()        

        console = logging.StreamHandler()    

        class IconFormatter(logging.Formatter):
            ICONS = {'DEBUG': '🔍', 'INFO': '✨', 'WARNING': '⚠️', 'ERROR': '❌', 'CRITICAL': '💥'}
            def format(self, record):
                icon = self.ICONS.get(record.levelname, 'ℹ️')
                record.icon = icon
                return super().format(record)       
            
        console.setFormatter(IconFormatter('%(icon)s %(message)s'))
        logger.addHandler(console)        

        file_handler = logging.FileHandler('crystal_evoluder.log', encoding='utf-8')
        file_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
        logger.addHandler(file_handler)        
        return logger

    def _check_grobid(self) -> bool:
        """Grobidサーバーが起動しているか確認"""
        try:
            response = requests.get(f"{self.grobid_url}/api/isalive", timeout=2)
            return response.status_code == 200
        except:
            return False

    def _init_grobid(self):
            """Grobidクライアント初期化"""
            from grobid_client.grobid_client import GrobidClient
        
            self.grobid_client = GrobidClient(
                grobid_server=self.config.get('grobid_url', 'http://localhost:8070'),
                timeout=120
            )
            self.logger.info("✅ Grobid client initialized")
    
    def _parse_pdf(self, pdf_path: Path) -> List[Document]:
        """PDF → GROBID → TEI → Document"""
        if not self.grobid_client:
            raise RuntimeError(
                f"Grobid server not available at {self.grobid_url}\n"
                "Start your Grobid server first"
)
        
        with self._safe_parse('pdf', pdf_path):
            # TEI出力先
            tei_path = pdf_path.with_suffix('.tei.xml')

            if tei_path.exists():
                self.logger.info(f"Using cached TEI: {tei_path.name}")
                return self._parse_tei(tei_path)

            # Grobid処理
            max_retries = 3
            for attempt in range(max_retries):    
                try:
                    self.logger.info(
                        f"Processing PDF with Grobid: {pdf_path.name}"
                        f"(attempt {attempt+1}/{max_retries})"
                    )
                
                    self.grobid_client.process(
                    'processFulltextDocument',
                    str(pdf_path),
                    output=str(tei_path.parent),
                    consolidate_citations=True,
                    consolidate_header=True,
                    include_raw_citations=False,
                    include_raw_affiliations=False,
                    tei_coordinates=False,
                    segment_sentences=False
                    )  
                    # 既存のTEIパーサーで処理
                    if tei_path.exists():
                        return self._parse_tei(tei_path)
                    else:
                        raise FileNotFoundError(f"Grobid output not found: {tei_path}")

                except requests.exceptions.Timeout:
                    self.logger.warning(f"⏱️  Timeout on attempt {attempt+1}")
                    if attempt == max_retries - 1:
                        raise RuntimeError("Grobid processing timeout")
                    time.sleep(2 ** attempt)  # exponential backoff

                except requests.exceptions.ConnectionError:
                    self.logger.error("❌ Grobid server unreachable")
                    raise RuntimeError(
                        "Grobid server not responding. Check:\n"
                        "1. Server is running: docker ps\n"
                        "2. URL is correct: http://localhost:8070\n"
                        "3. Restart: docker restart <container_id>"
                    )

                except Exception as e:
                    self.logger.error(f"❌ Grobid processing failed: {e}")
                    if attempt == max_retries - 1:
                        raise
                    time.sleep(2 ** attempt)

    def batch_crystallize(self, input_dir: str, 
                            patterns: List[str] = None,
                            max_workers: int = 4,
                            fail_fast: bool = False) -> Dict[str, Any]:

            """
            ディレクトリ内のファイルを一括処理
        
            Args:
                include_pdf: Grobid有効時のみPDFを処理（None=自動判定）
                patterns: ファイルパターン（例: ['*.md', '*.pdf']）
                max_workers: 並列処理数
                fail_fast: True=最初のエラーで停止、False=全部試す
            Returns:
                {
                'success': {filepath: [Document, ...], ...},
                'failed': [(filepath, error_msg), ...],
                'skipped': [filepath, ...],  # PDF skip等
                'stats': {...}
                }
            """
            if include_pdf is None:
                include_pdf = self.grobid_available

            if patterns is None:
                patterns = patterns or ['*.md', '*.docx', '*.html', '*.txt', '*.tei.xml']
        
                # Grobid有効時のみPDFを追加
                if self.grobid_available:
                    patterns.append('*.pdf')
                    self.logger.info("✅ PDF processing enabled")
                else:
                    self.logger.warning("⚠️  PDF skipped (Grobid server not available)")
            # ファイル収集
            files = []        
            for pattern in patterns:
                files.extend(Path(input_dir).rglob(pattern))
        
            self.logger.info(f"Found {len(files)} files")
             # 処理結果
            results = {
                'success': {},
                'failed': [],
                'skipped': [],
                'stats': {}
            }
            
            # 並列処理
            with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
                future_to_file = {} 

                for f in files:
                    if f.suffix.lower() == '.pdf' and not include_pdf:
                        results['skipped'].append(str(f))
                        continue    
                    future = executor.submit(self._crystallize_with_retry, str(f))
                    future_to_file[future] = f

                # プログレスバー
                try:
                    from tqdm import tqdm
                    iterator = tqdm(
                
                        concurrent.futures.as_completed(future_to_file),
                        total=len(future_to_file),
                        desc="Crystallizing"
                    )
                except ImportError:
                    iterator = concurrent.futures.as_completed(future_to_file)
                # 結果収集
                for future in iterator:
                    file_path = future_to_file[future]
                    try:
                        docs = future.result(timeout=300)
                        results[str(file_path)] = docs
                    
                    except concurrent.futures.TimeoutError:
                        error_msg = "Processing timeout (>5min)"
                        results['failed'].append((str(file_path), error_msg))
                        self.logger.error(f"Failed: {file_path.name} : {error_msg}")
                        if fail_fast:
                            executor.shutdown(wait=False)
                            break

                    except Exception as e:
                        error_msg = str(e)
                        results['failed'].append((str(file_path), error_msg))
                        self.logger.error(f"Failed: {file_path.name} : {error_msg}")
                        if fail_fast:
                            executor.shutdown(wait=False)
                            break
        # 統計情報
            results['stats'] = {
                'total': len(files),
                'success': len(results['success']),
                'failed': len(results['failed']),
                'skipped': len(results['skipped']),
                'total_documents': sum(len(docs) for docs in results['success'].values())
        }            
        
            # サマリー
            self.logger.info(
                f"✅ Batch complete: "
                f"{results['stats']['success']} success, "
                f"{results['stats']['failed']} failed, "
                f"{results['stats']['skipped']} skipped"
            )
        
        # 失敗レポート保存
            if results['failed']:
                self._save_error_report(results['failed'], input_dir)
        
            return results

    def _save_error_report(self, failed: List[Tuple[str, str]], base_dir: str):
        """エラーレポート保存"""
        report_path = Path(base_dir) / 'crystal_evoluder_errors.json'
    
        report = {
            'timestamp': datetime.now().isoformat(),
            'failed_files': [
                {'file': filepath, 'error': error}
                for filepath, error in failed
            ]
        }
    
        with open(report_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, indent=2, ensure_ascii=False)
    
        self.logger.info(f"📝 Error report saved: {report_path}")


    def _crystallize_with_retry(self, file_path: str, 
                                max_retries: int = 3) -> List[Document]:
        """リトライ機構付き crystallize"""
        for attempt in range(max_retries):
            try:
                return self.crystallize(file_path)
            except Exception as e:
                if attempt == max_retries - 1:
                    raise
                self.logger.warning(f"Retry {attempt+1}/{max_retries}: {file_path}\n"
                                    f"  Error: {e}")
                time.sleep(2 ** attempt)  # exponential backoff

    def crystallize(self, input_path: str, format: str = 'auto') -> List[Document]:
        """結晶化: 入力ファイルをパース"""
        self.logger.info("Crystallizing knowledge structure...")
        
        input_path = Path(input_path).expanduser()
        
        if not input_path.exists():
            raise FileNotFoundError(f"File not found: {input_path}")
        
        if format == 'auto':
            format = self._detect_format(str(input_path))
        
        parsers = {
            'tei': self._parse_tei,
            'markdown': self._parse_markdown,
            'txt': self._parse_txt,
            'docx': self._parse_docx,
            'html': self._parse_html
        }
        
        if format not in parsers:
            raise ValueError(f"Unsupported format: {format}")
        
        documents = parsers[format](input_path)
        
        self.crystal = documents
        self.logger.info(f"✨ Crystal structure stabilized: {len(documents)} nodes")
        
        return documents
    
    def _detect_format(self, input_path: str) -> str:
        """ファイル形式を自動判定"""
        path = Path(input_path)
        suffixes = ''.join(path.suffixes).lower()
        
        if suffixes.endswith('.tei.xml'):
            return 'tei'
        
        suffix = path.suffix.lower()
        format_map = {
            '.xml': 'xml',
            '.pdf': 'pdf',
            '.md': 'markdown',
            '.markdown': 'markdown',
            '.docx': 'docx',
            '.txt': 'txt',
            '.html': 'html',
            '.htm': 'html'
        }
        
        detected = format_map.get(suffix)
        
        if detected == 'xml' and self._is_tei_xml(path):
            detected = 'tei'
        
        if not detected:
            raise ValueError(f"Unsupported format: {suffix}")
        
        self.logger.info(f"Detected format: {detected}")
        return detected
    
    def _is_tei_xml(self, xml_path: Path) -> bool:
        """XMLがTEIか判定"""
        try:
            with open(xml_path, 'r', encoding='utf-8') as f:
                header = f.read(1000)
                return '<teiHeader>' in header or '<TEI' in header
        except:
            return False
    
    @contextmanager
    def _safe_parse(self, format_name: str, input_path: Path):
        """パーサーの例外処理"""
        try:
            yield
        except Exception as e:
            self.logger.error(f"[{format_name}] Parse failed: {input_path}\n{e}", exc_info=True)
            self.crystal = []
            self.metadata = {'title': input_path.stem, 'authors': ['Unknown']}
    
    def _clean_text(self, text: str) -> str:
        """テキスト正規化"""
        if HAS_FTFY:
            text = ftfy.fix_text(text)
        text = re.sub(r'\s+', ' ', text).strip()
        return text
    
    def _parse_tei(self, tei_path: Path) -> List[Document]:
        """TEIをパース"""
        with self._safe_parse('tei', tei_path):
            with open(tei_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f, 'xml')
            
            title = self._extract_title_safe(soup, tei_path)
            authors = self._extract_authors_safe(soup)
            
            documents = []
            for i, div in enumerate(soup.find_all('div')):
                doc = self._extract_section_safe(div, i, title, authors)
                if doc:
                    documents.append(doc)
            
            self.metadata = {'title': title, 'authors': authors}
            self.logger.info(f"TEI parsed: {len(documents)} sections")
            return documents
    
    def _extract_title_safe(self, soup: BeautifulSoup, filepath: Path) -> str:
        """タイトル抽出"""
        try:
            title_tag = soup.find('titleStmt')
            if title_tag:
                title_tag = title_tag.find('title', level='a', type='main')
            title = title_tag.text.strip() if title_tag else None
            if title and len(title) > 10:
                return title
        except:
            pass
        return filepath.stem.replace('_', ' ').title()
    
    def _extract_authors_safe(self, soup: BeautifulSoup) -> List[str]:
        """著者抽出"""
        authors = []
        try:
            for persName in soup.find_all('persName'):
                try:
                    forenames = [f.text for f in persName.find_all('forename') if f.text]
                    surname = persName.find('surname')
                    author_name = f"{' '.join(forenames)} {surname.text if surname else ''}".strip()
                    if author_name:
                        authors.append(author_name)
                except:
                    continue
        except:
            pass
        return authors if authors else ["Unknown Author"]
    
    def _extract_section_safe(self, div, index: int, title: str, authors: List[str]) -> Optional[Document]:
        """セクション抽出"""
        try:
            head = div.find('head')
            section_title = head.text.strip() if head and head.text else f"Section {index}"
            
            paragraphs = [p.get_text(strip=True) for p in div.find_all('p') if p.get_text(strip=True)]
            text = '\n\n'.join(paragraphs)
            
            if len(text) < 50:
                return None
            
            return Document(
                text=self._clean_text(text),
                metadata={
                    'title': title[:200],
                    'authors': ', '.join(authors[:5]),
                    'section': section_title[:100],
                    'section_index': index,
                    'source_format': 'tei'
                }
            )
        except:
            return None
    
    def _parse_markdown(self, md_path: Path) -> List[Document]:
        """Markdownをパース"""
        with self._safe_parse('markdown', md_path):
            with open(md_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if HAS_FRONTMATTER:
                post = frontmatter.load(md_path)
                content = post.content
                title = post.metadata.get('title', md_path.stem)
                authors = post.metadata.get('authors', ['Unknown'])
            else:
                title = md_path.stem
                authors = ['Unknown']
            
            sections = self._split_by_headers_safe(content)
            
            documents = []
            for i, (heading, text) in enumerate(sections):
                clean_text = self._clean_text(text)
                if len(clean_text) >= 50:
                    documents.append(Document(
                        text=clean_text,
                        metadata={
                            'title': title,
                            'authors': ', '.join(authors) if isinstance(authors, list) else authors,
                            'section': heading,
                            'section_index': i,
                            'source_format': 'markdown'
                        }
                    ))
            
            self.metadata = {'title': title, 'authors': authors if isinstance(authors, list) else [authors]}
            self.logger.info(f"Markdown parsed: {len(documents)} sections")
            return documents
    
    def _split_by_headers_safe(self, content: str) -> List[Tuple[str, str]]:
        """Markdownを見出しで分割（コードブロック対応）"""
        code_blocks = []
        def replace_code(match):
            code_blocks.append(match.group(0))
            return f"__CODE_BLOCK_{len(code_blocks)-1}__"
        
        content_no_code = re.sub(r'```.*?```', replace_code, content, flags=re.DOTALL)
        
        pattern = r'^(#{1,6})\s+(.+)$'
        lines = content_no_code.split('\n')
        
        sections = []
        current_heading = "Introduction"
        current_text = []
        
        for line in lines:
            match = re.match(pattern, line)
            if match:
                if current_text:
                    text = '\n'.join(current_text)
                    for i, block in enumerate(code_blocks):
                        text = text.replace(f"__CODE_BLOCK_{i}__", block)
                    sections.append((current_heading, text))
                current_heading = match.group(2)
                current_text = []
            else:
                current_text.append(line)
        
        if current_text:
            text = '\n'.join(current_text)
            for i, block in enumerate(code_blocks):
                text = text.replace(f"__CODE_BLOCK_{i}__", block)
            sections.append((current_heading, text))
        
        return sections
    
    def _parse_txt(self, txt_path: Path) -> List[Document]:
        """プレーンテキストをパース"""
        with self._safe_parse('txt', txt_path):
            with open(txt_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            content = self._clean_text(content)

        # APIキーの取得（優先順位）
        api_key = (
            self.config.get('openai_api_key') or  # 1. 設定ファイル/引数
            os.environ.get('OPENAI_API_KEY')      # 2. 環境変数（.zshrc）
        )

        # セマンティックチャンキング
        if HAS_SEMANTIC_CHUNKER and api_key:
            try:
                os.environ['OPENAI_API_KEY'] = api_key
                
                self.logger.info("🤖 Using SemanticChunker")
                embeddings = OpenAIEmbeddings()
                splitter = SemanticChunker(embeddings, breakpoint_threshold_type="percentile")
                chunks = splitter.split_text(content)
                self.logger.info(f"✅ {len(chunks)} semantic chunks created")
            except Exception as e:
                self.logger.warning(f"⚠️  SemanticChunker failed: {e}")
                chunks = self._chunk_by_paragraphs(content)
        else:
            if HAS_SEMANTIC_CHUNKER and not self.config.get('openai_api_key'):
                self.logger.info("ℹ️  OpenAI API key not provided, using basic chunking")
            chunks = self._chunk_by_paragraphs(content)
            
            documents = []
            for i, chunk in enumerate(chunks):
                documents.append(Document(
                    text=chunk,
                    metadata={
                        'title': txt_path.stem,
                        'authors': 'Unknown',
                        'section': f"Chunk {i+1}",
                        'section_index': i,
                        'source_format': 'txt'
                    }
                ))
            
            self.metadata = {'title': txt_path.stem, 'authors': ['Unknown']}
            self.logger.info(f"Text parsed: {len(documents)} chunks")
            return documents
    
    def _chunk_by_paragraphs(self, content: str, chunk_size: int = 2000) -> List[str]:
        """段落ベースのチャンキング"""
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        chunks = []
        current_chunk = []
        current_length = 0
        
        for para in paragraphs:
            para_length = len(para)
            if current_length + para_length > chunk_size and current_chunk:
                chunks.append('\n\n'.join(current_chunk))
                current_chunk = []
                current_length = 0
            current_chunk.append(para)
            current_length += para_length
        
        if current_chunk:
            chunks.append('\n\n'.join(current_chunk))
        
        return chunks
    
    def _parse_docx(self, docx_path: Path) -> List[Document]:
        """Word文書をパース"""
        if not HAS_DOCX:
            raise ImportError("python-docx not installed")
        
        with self._safe_parse('docx', docx_path):
            doc = DocxDocument(docx_path)
            title = doc.core_properties.title or docx_path.stem
            authors = doc.core_properties.author or 'Unknown'
            
            heading_styles = {'heading 1', 'heading 2', 'heading 3', 'heading1', 'heading2', 'heading3', '見出し 1', '見出し 2', '見出し 3', 'title'}
            
            sections = []
            current_heading = "Introduction"
            current_text = []
            
            for para in doc.paragraphs:
                is_heading = para.style.name.lower() in heading_styles
                if is_heading:
                    if current_text:
                        sections.append((current_heading, self._clean_text('\n\n'.join(current_text))))
                    current_heading = para.text
                    current_text = []
                else:
                    if para.text.strip():
                        current_text.append(para.text)
            
            if current_text:
                sections.append((current_heading, self._clean_text('\n\n'.join(current_text))))
            
            documents = []
            for i, (heading, text) in enumerate(sections):
                if len(text) >= 50:
                    documents.append(Document(
                        text=text,
                        metadata={'title': title, 'authors': authors, 'section': heading, 'section_index': i, 'source_format': 'docx'}
                    ))
            
            self.metadata = {'title': title, 'authors': [authors] if isinstance(authors, str) else authors}
            self.logger.info(f"DOCX parsed: {len(documents)} sections")
            return documents

    def _parse_html(self, html_path: Path) -> List[Document]:
        """HTMLをパース"""
        with self._safe_parse('html', html_path):
            with open(html_path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            # 1) trafilatura を試す（整形済みテキストが欲しい時）
            if HAS_TRAFILATURA:
                extracted = trafilatura.extract(html_content, include_formatting=True)                
                if extracted:
                    text = self._clean_text(extracted)
                    title = html_path.stem

                    soup = BeautifulSoup(html_content, 'html.parser')
                    authors = self._extract_html_authors(soup)

                    self.metadata = {'title': title, 'authors': authors}
                    self.logger.info("HTML parsed with trafilatura")
                    return [Document(
                        text=text, 
                        metadata={'title': title, 'authors': ', '.join(authors),
                        'section': 'Main',
                        'source_format': 'html'}
                    )]

            # 2) fallback BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')
            title_tag = soup.find('title')
            title = title_tag.get_text(strip=True) if title_tag else html_path.stem

            # 著者メタ情報取得
            authors = self._extract_html_authors(soup) 

            # 見出しと段落でセクション分割
            sections = []
            current_heading = "Introduction"
            current_text = []        

            for element in soup.find_all(['h1', 'h2', 'h3', 'p']):
                if element.name in ['h1', 'h2', 'h3']:
                    if current_text:
                        sections.append((current_heading, self._clean_text('\n\n'.join(current_text))))
                    current_heading = element.get_text(strip=True)
                    current_text = []
                elif element.name == 'p':
                    text = element.get_text(strip=True)
                    if text:
                        current_text.append(text)

            if current_text:
                sections.append((current_heading, self._clean_text('\n\n'.join(current_text))))

            documents = []
            for i, (heading, text) in enumerate(sections):
                if len(text) >= 50:
                    documents.append(Document(
                        text=text,
                        metadata={
                            'title': title,
                            'authors': ', '.join(authors),
                            'section': heading,
                            'section_index': i,
                            'source_format': 'html'
                        }
                    ))

            self.metadata = {'title': title, 'authors': authors}
            self.logger.info(f"HTML parsed: {len(documents)} sections")
            return documents

    def _extract_html_authors(self, soup: BeautifulSoup) -> List[str]:
        """HTML から著者情報を抽出（共通メソッド）"""
         # 複数のメタタグパターンを試す
        for meta_name in ('author', 'article:author', 'og:author', 'byline'):
            meta = soup.find('meta', attrs={'name': meta_name})
            if not meta:
                meta = soup.find('meta', attrs={'property': meta_name})
        
            if meta and meta.get('content'):
                author = meta.get('content').strip()
                if author:
                    return [author]
    
        # rel="author" を試す
        author_link = soup.find('a', rel='author')
        if author_link:
            author = author_link.get_text(strip=True)
            if author:
                return [author]
    
    # 見つからなければ Unknown
        return ['Unknown']

    def evolve_to_notes(self, output_dir: str):
        """ノートに進化"""
        if not self.crystal:
            raise ValueError("Run crystallize() first")
        
        self.logger.info("Evolving to notes...")
        output_dir = Path(output_dir).expanduser()
        paper_title = self._sanitize(self.metadata['title'])
        paper_dir = output_dir / "Papers" / paper_title
        paper_dir.mkdir(parents=True, exist_ok=True)
        
        for i, doc in enumerate(self.crystal):
            section = doc.metadata.get('section', 'Untitled')
            md_content = f"""---
title: {self.metadata['title']}
authors: {', '.join(self.metadata['authors'][:3])}
section: {section}
---

# {section}

{doc.text}
"""
            filename = f"{paper_title}_{i:03d}_{self._sanitize(section)}.md"
            (paper_dir / filename).write_text(md_content, encoding='utf-8')
        
        self.logger.info(f"✅ Notes created: {paper_dir}")
    
    def evolve_to_graph(self, graph_store: Neo4jGraphStore):
        """Knowledge Graphに進化"""
        if not self.crystal:
            raise ValueError("Run crystallize() first")
        
        self.logger.info("Evolving to graph...")
        llm = OpenAI(model="gpt-4o-mini", timeout=120.0)
        storage_context = StorageContext.from_defaults(graph_store=graph_store)
        
        index = KnowledgeGraphIndex.from_documents(
            self.crystal,
            storage_context=storage_context,
            llm=llm,
            transformations=[SimpleNodeParser.from_defaults(chunk_size=512)],
            embed_model=HuggingFaceEmbedding(model_name="BAAI/bge-m3"),
            show_progress=False,
            max_triplets_per_chunk=10
        )
        
        kg = index.get_networkx_graph()
        self.logger.info(f"✅ Graph created: {len(kg.nodes)} nodes, {len(kg.edges)} edges")
    
    def _sanitize(self, text: str) -> str:
        """ファイル名用サニタイズ"""
        return re.sub(r'[<>:"/\\|?*]', '', text)[:50].strip()


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description='Crystal Evoluder v1.1')
    parser.add_argument('command', choices=[
        'crystallize', 
        'evolve-notes', 
        'evolve-graph', 
        'evolve-all',
        'batch'  
   ])

    parser.add_argument('input_file', help='Input file or directory path')
    parser.add_argument('--format', default='auto', help='File format (auto/tei/markdown/docx/html/txt/pdf)')
    parser.add_argument('--markdown-dir', default='~/CrystalEvoluder/Library')
    parser.add_argument('--neo4j-uri', default='bolt://localhost:7687')
    parser.add_argument('--neo4j-user', default='neo4j')
    parser.add_argument('--neo4j-pass')
    parser.add_argument('--debug', action='store_true')
    parser.add_argument('--max-workers', type=int, default=4, help='Parallel workers for batch')
    parser.add_argument('--fail-fast', action='store_true', help='Stop on first error')

    args = parser.parse_args()
    
    print("🔮 Crystal Evoluder v1.1")
    print("━" * 42)
    
    evoluder = CrystalEvoluder(log_level=logging.DEBUG if args.debug else logging.INFO)
    evoluder.crystallize(args.input_file, format=args.format)
    
    if args.command == 'crystallize':
        evoluder.crystallize(args.input_file, format=args.format)

    if args.command == 'batch':
        results = evoluder.batch_crystallize(
            args.input_file,
            max_workers=args.max_workers,
            fail_fast=args.fail_fast
        )
        
        # 結果サマリー表示
        print("\n" + "="*42)
        print("📊 Batch Processing Results")
        print("="*42)
        print(f"✅ Success: {results['stats']['success']}")
        print(f"❌ Failed:  {results['stats']['failed']}")
        print(f"⏭️  Skipped: {results['stats']['skipped']}")
        print(f"📄 Total documents: {results['stats']['total_documents']}")
        
        if results['failed']:
            print("\n❌ Failed files:")
            for filepath, error in results['failed'][:5]:  # 最初の5件
                print(f"  - {Path(filepath).name}: {error[:50]}...")

    if args.command in ['evolve-notes', 'evolve-all']:
        evoluder.evolve_to_notes(args.markdown_dir)
    
    if args.command in ['evolve-graph', 'evolve-all']:
        if not args.neo4j_pass:
            raise ValueError("--neo4j-pass required")
        graph_store = Neo4jGraphStore(username=args.neo4j_user, password=args.neo4j_pass, url=args.neo4j_uri)
        evoluder.evolve_to_graph(graph_store)

    print("✨ Complete!")