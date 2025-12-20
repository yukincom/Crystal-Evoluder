"""
DOCX Parser
Word文書のパース
"""
from pathlib import Path
from typing import List
from llama_index.core import Document

from ..text_utils import clean_text

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

def parse_docx(docx_path: Path, logger=None) -> List[Document]:
    """Word文書をパース"""
    if not HAS_DOCX:
        if logger:
            logger.error("python-docx not installed")
        raise ImportError("python-docx not installed")

    try:
        doc = DocxDocument(docx_path)
        title = doc.core_properties.title or docx_path.stem
        authors = doc.core_properties.author or 'Unknown'

        heading_styles = {'heading 1', 'heading 2', 'heading 3', 'heading1', 'heading2', 'heading3', '見出し 1', '見出し 2', '見出し 3', 'title'}

        sections = []
        current_heading = "Introduction"
        current_text = []

        for para in doc.paragraphs:
            is_heading = para.style.name.lower() in heading_styles
            if is_heading:
                if current_text:
                    sections.append((current_heading, clean_text('\n\n'.join(current_text))))
                current_heading = para.text
                current_text = []
            else:
                if para.text.strip():
                    current_text.append(para.text)

        if current_text:
            sections.append((current_heading, clean_text('\n\n'.join(current_text))))

        documents = []
        for i, (heading, text) in enumerate(sections):
            if len(text) >= 50:
                documents.append(Document(
                    text=text,
                    metadata={'title': title, 'authors': authors, 'section': heading, 'section_index': i, 'source_format': 'docx'}
                ))

        metadata = {'title': title, 'authors': [authors] if isinstance(authors, str) else authors}
        if logger:
            logger.info(f"DOCX parsed: {len(documents)} sections")
        return documents, metadata

    except Exception as e:
        if logger:
            logger.error(f"DOCX parse failed: {e}")
        return [], {'title': docx_path.stem, 'authors': ['Unknown']}